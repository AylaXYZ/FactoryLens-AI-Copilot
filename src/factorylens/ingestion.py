from __future__ import annotations

import hashlib
from collections.abc import Iterable
from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md"}


def _clean(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


def load_document(path: str | Path) -> list[Document]:
    """Parse common enterprise files while preserving source/page/sheet metadata."""
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")

    base = {"source": file_path.name, "path": str(file_path)}
    documents: list[Document] = []

    if suffix == ".pdf":
        for page_no, page in enumerate(PdfReader(str(file_path)).pages, start=1):
            text = _clean(page.extract_text() or "")
            if text:
                documents.append(Document(page_content=text, metadata={**base, "page": page_no}))
    elif suffix == ".docx":
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = _clean("\n".join(paragraphs))
        if text:
            documents.append(Document(page_content=text, metadata=base))
    elif suffix in {".xlsx", ".xls"}:
        sheets = pd.read_excel(file_path, sheet_name=None)
        for sheet_name, frame in sheets.items():
            text = frame.fillna("").astype(str).to_csv(index=False)
            if text.strip():
                documents.append(
                    Document(page_content=text, metadata={**base, "sheet": str(sheet_name)})
                )
    elif suffix == ".csv":
        frame = pd.read_csv(file_path)
        documents.append(
            Document(page_content=frame.fillna("").astype(str).to_csv(index=False), metadata=base)
        )
    else:
        text = _clean(file_path.read_text(encoding="utf-8"))
        if text:
            documents.append(Document(page_content=text, metadata=base))

    return documents


def chunk_documents(
    documents: Iterable[Document], chunk_size: int = 700, chunk_overlap: int = 120
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n## ", "\n### ", "\n", "。", "；", " ", ""],
    )
    chunks = splitter.split_documents(list(documents))
    for chunk in chunks:
        raw = f"{chunk.metadata.get('source')}::{chunk.page_content}".encode()
        chunk.metadata["chunk_id"] = hashlib.sha1(raw).hexdigest()[:12]
    return chunks


def load_and_chunk(paths: Iterable[str | Path]) -> list[Document]:
    documents: list[Document] = []
    for path in paths:
        documents.extend(load_document(path))
    return chunk_documents(documents)

