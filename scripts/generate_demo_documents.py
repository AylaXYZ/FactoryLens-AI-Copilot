"""Generate synthetic Word, Excel and PDF files to demonstrate the ingestion pipeline."""

from pathlib import Path

import pandas as pd
from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "demo_documents"
OUTPUT.mkdir(exist_ok=True)


def generate_docx() -> None:
    document = Document()
    document.add_heading("设备点检作业指导书（演示）", level=1)
    document.add_paragraph("每日开机前检查风扇、联轴器、紧固件和异常噪声。")
    document.add_paragraph("发现温度超过 75°C 时，记录设备编号并通知设备工程师。")
    document.save(OUTPUT / "设备点检作业指导书.docx")


def generate_xlsx() -> None:
    frame = pd.DataFrame(
        [
            {"备件编码": "SP-FAN-01", "名称": "冷却风扇", "安全库存": 2, "当前库存": 4},
            {"备件编码": "SP-BRG-07", "名称": "驱动端轴承", "安全库存": 3, "当前库存": 1},
        ]
    )
    frame.to_excel(OUTPUT / "关键备件台账.xlsx", index=False, sheet_name="备件")


def generate_pdf() -> None:
    # Uses a Windows CJK font when available; falls back to an English-only demo.
    output = OUTPUT / "维修安全须知.pdf"
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    page = canvas.Canvas(str(output))
    if font_path.exists():
        pdfmetrics.registerFont(TTFont("CJK", str(font_path), subfontIndex=0))
        page.setFont("CJK", 14)
        lines = ["维修安全须知（演示）", "维修前必须停机、断电并挂牌上锁。", "AI 建议必须由工程师复核。"]
    else:
        page.setFont("Helvetica", 14)
        lines = ["Maintenance Safety Demo", "Lock out and tag out before maintenance."]
    y = 780
    for line in lines:
        page.drawString(72, y, line)
        y -= 28
    page.save()


if __name__ == "__main__":
    generate_docx()
    generate_xlsx()
    generate_pdf()
    print(f"Generated demo documents in {OUTPUT}")

